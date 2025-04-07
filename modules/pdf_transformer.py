# modules/pdf_transformer.py

import os
import tempfile
from pathlib import Path
import logging
import platform
import shutil
import subprocess
import io # Para trabalhar com bytes em memória
import zipfile
from PyPDF2 import PdfReader, PdfWriter # Import mantido para merge_pdfs

try:
    import fitz  # PyMuPDF
except ImportError:
    logging.warning("PyMuPDF (fitz) não encontrado. Funções de processamento PDF avançadas estarão indisponíveis.")
    fitz = None

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    logging.warning("Módulo python-docx não encontrado. Funções de conversão PDF para DOCX estarão indisponíveis.")
    Document = None

try:
    import img2pdf
    from PIL import Image
except ImportError:
    logging.warning("Módulos img2pdf ou Pillow não encontrados. Funções de conversão de imagem estarão indisponíveis.")
    img2pdf = None
    Image = None

# Configuração de Logging (pode ser configurado centralmente em app.py também)
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
# Logger nomeado para este módulo
log = logging.getLogger(__name__)

class PDFTransformer:
    def __init__(self):
        self.ocrmypdf_cmd = 'ocrmypdf'
        # Usar shutil.which para encontrar executáveis de forma mais confiável
        self.libreoffice_path = self._find_external_command(['libreoffice', 'soffice'], 'LibreOffice')
        self.gs_cmd = self._find_external_command(['gswin64c', 'gs'], 'Ghostscript') # Encontra gs ou gswin64c
        self.ocrmypdf_installed = self._check_command_exists(self.ocrmypdf_cmd, 'OCRmyPDF')

        # Log inicial sobre dependências
        if not self.gs_cmd: log.warning("Ghostscript não detectado. Funções de compressão/manipulação podem ser limitadas.")
        if not self.ocrmypdf_installed: log.warning("OCRmyPDF não detectado. Funções de OCR estarão indisponíveis.")
        if not self.libreoffice_path: log.warning("LibreOffice não detectado. Funções de conversão de/para DOC/XLS/ODS estarão indisponíveis.")

    def _find_external_command(self, potential_cmds, tool_name):
        """Tenta encontrar um comando em uma lista de nomes possíveis."""
        for cmd in potential_cmds:
            found_path = shutil.which(cmd)
            if found_path:
                log.info(f"{tool_name} detectado em: {found_path}")
                return found_path # Retorna o caminho completo encontrado
        log.warning(f"{tool_name} (comandos testados: {', '.join(potential_cmds)}) não encontrado no PATH.")
        return None

    def _check_command_exists(self, cmd, tool_name):
        """Verifica se um comando específico existe e é executável."""
        if shutil.which(cmd):
            # Pode adicionar uma verificação de versão mais robusta se necessário
            # try:
            #     subprocess.run([cmd, '--version'], capture_output=True, check=True, timeout=5)
            #     log.info(f"{tool_name} detectado e funcional.")
            #     return True
            # except (subprocess.SubprocessError, FileNotFoundError, Exception) as e:
            #     log.warning(f"Erro ao verificar {tool_name} com '--version': {e}")
            #     return False # Existe mas pode não funcionar
            log.info(f"{tool_name} (comando: {cmd}) parece estar disponível no PATH.")
            return True # Assume que existe se shutil.which encontrar
        else:
            # Log já feito em _find_external_command se aplicável, mas pode repetir para clareza
            # log.warning(f"{tool_name} (comando: {cmd}) não encontrado no PATH.")
            return False

    def _run_subprocess(self, command, description, timeout=300):
        """Executa um subprocesso com logging e tratamento de erro."""
        log.info(f"Executando {description}: {' '.join(command)}")
        try:
            process = subprocess.run(command, capture_output=True, check=False, text=True, timeout=timeout)
            if process.returncode != 0:
                log.error(f"Erro em {description} (código {process.returncode}):\nStderr: {process.stderr.strip()}\nStdout: {process.stdout.strip()}")
                return False, f"Erro durante {description} (código {process.returncode})."
            log.info(f"{description} concluído com sucesso.")
            return True, process # Retorna o processo para análise posterior se necessário
        except subprocess.TimeoutExpired:
            log.error(f"{description} excedeu o tempo limite de {timeout}s.")
            return False, f"{description.capitalize()} excedeu o tempo limite."
        except FileNotFoundError:
            log.error(f"Comando não encontrado para {description}: {command[0]}")
            return False, f"Comando necessário para {description} não encontrado."
        except Exception as e:
            log.exception(f"Erro inesperado ao executar {description}: {e}") # Loga traceback
            return False, f"Erro inesperado durante {description}."

    # --- Métodos Principais (Lógica interna praticamente inalterada, apenas logging ajustado) ---

    def _apply_ocrmypdf(self, input_path, output_path, language='por'):
        if not self.ocrmypdf_installed:
            log.error("Tentativa de usar OCRmyPDF, mas não está instalado/detectado.")
            return False, "OCRmyPDF não está disponível no servidor."
        # Usar diretório temporário gerenciado
        with tempfile.TemporaryDirectory() as temp_ocr_dir:
            # Copiar entrada para diretório temporário para evitar problemas de permissão/nomes
            temp_input = Path(temp_ocr_dir) / "input_ocr.pdf"
            try:
                shutil.copy2(input_path, temp_input)
            except Exception as copy_err:
                log.error(f"Erro ao copiar arquivo para OCR: {copy_err}")
                return False, "Erro ao preparar arquivo para OCR."

            temp_output = Path(temp_ocr_dir) / "output_ocr.pdf"
            args = [
                self.ocrmypdf_cmd, '--force-ocr', '--optimize', '1', '--output-type', 'pdf',
                '--jobs', '2', '-l', language, str(temp_input), str(temp_output)
            ]
            # Usar o helper _run_subprocess
            success, result_msg_or_proc = self._run_subprocess(args, "OCRmyPDF", timeout=600) # Timeout maior para OCR

            if not success:
                return False, result_msg_or_proc # Retorna a mensagem de erro do helper

            # Verificar se o arquivo de saída foi criado
            if temp_output.exists() and temp_output.stat().st_size > 0:
                try:
                    # Mover o resultado para o caminho de saída final
                    shutil.move(str(temp_output), output_path)
                    log.info(f"OCR aplicado com sucesso e salvo em {output_path}")
                    return True, "OCR aplicado com sucesso."
                except Exception as move_err:
                    log.error(f"OCRmyPDF executado, mas erro ao mover o resultado: {move_err}")
                    return False, "Erro ao salvar o resultado do OCR."
            else:
                log.error("OCRmyPDF terminou sem erros aparentes, mas o arquivo de saída não foi encontrado ou está vazio.")
                # Incluir stdout/stderr se possível (result_msg_or_proc é o 'process' object aqui)
                if isinstance(result_msg_or_proc, subprocess.CompletedProcess):
                    log.error(f"Stderr: {result_msg_or_proc.stderr.strip()}\nStdout: {result_msg_or_proc.stdout.strip()}")
                return False, "Falha na criação do arquivo OCRizado."

    def _compress_pdf_gs(self, input_file_path, output_file_path, power=3):
        if not self.gs_cmd:
            log.error("Ghostscript não está disponível. Não é possível comprimir.")
            return False, "Ghostscript não encontrado no servidor."

        quality = {0: '/default', 1: '/prepress', 2: '/printer', 3: '/ebook', 4: '/screen'}
        power = max(0, min(4, power)) # Garante que power está entre 0 e 4
        command = [
            self.gs_cmd, '-sDEVICE=pdfwrite', '-dCompatibilityLevel=1.4',
            f'-dPDFSETTINGS={quality[power]}', '-dNOPAUSE', '-dQUIET', '-dBATCH',
            f'-sOutputFile={output_file_path}', input_file_path
        ]

        # Usar o helper _run_subprocess
        success, result_msg_or_proc = self._run_subprocess(command, "Compressão Ghostscript", timeout=600) # Timeout maior

        if not success:
            # Tentar remover arquivo de saída parcial se existir
            if os.path.exists(output_file_path):
                try: os.unlink(output_file_path)
                except OSError as e: log.warning(f"Não foi possível remover arquivo de saída parcial (GS Error) {output_file_path}: {e}")
            return False, result_msg_or_proc # Mensagem de erro do helper

        # Verificar se o arquivo de saída foi criado e tem conteúdo
        if os.path.exists(output_file_path) and os.path.getsize(output_file_path) > 0:
            log.info(f"Compressão com Ghostscript concluída: {os.path.basename(output_file_path)}")
            return True, "Compressão bem-sucedida."
        else:
            log.error("Ghostscript finalizou sem erro aparente, mas o arquivo de saída não foi criado ou está vazio.")
            # Tentar remover arquivo de saída vazio se existir
            if os.path.exists(output_file_path):
                try: os.unlink(output_file_path)
                except OSError as e: log.warning(f"Não foi possível remover arquivo de saída vazio (GS) {output_file_path}: {e}")
            return False, "Falha na criação do arquivo comprimido."

    # Método principal para compressão e/ou OCR
    def process_compression_ocr(self, file_bytes, compression_level=3, apply_ocr=False, ocr_language='por'):
        """
        Processa um PDF aplicando compressão e/ou OCR.

        Args:
            file_bytes (bytes): Conteúdo do PDF original.
            compression_level (int): Nível de compressão Ghostscript (-1 para pular, 0-4).
            apply_ocr (bool): Se True, aplica OCR com OCRmyPDF.
            ocr_language (str): Código do idioma para OCR (ex: 'por', 'eng').

        Returns:
            tuple: (bool: success, bytes | None: processed_bytes, str: message)
        """
        if not file_bytes:
            return False, None, "Nenhum dado de arquivo fornecido."

        original_size_mb = len(file_bytes) / 1024 / 1024
        log.info(f"Iniciando processamento (Compressão: {compression_level}, OCR: {apply_ocr}, Lang: {ocr_language}). Tamanho Original: {original_size_mb:.2f} MB")

        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "input.pdf"
            with open(input_path, "wb") as f: f.write(file_bytes)

            current_step_output = input_path
            last_successful_output = input_path # Guarda o último resultado válido
            step_message = "Processo iniciado."

            # Etapa 1: Compressão (se solicitada)
            if compression_level >= 0:
                compressed_path = Path(temp_dir) / "compressed.pdf"
                log.info(f"Aplicando compressão (Nível {compression_level})...")
                comp_success, comp_msg = self._compress_pdf_gs(str(current_step_output), str(compressed_path), compression_level)
                step_message = comp_msg
                if comp_success:
                    current_step_output = compressed_path
                    last_successful_output = compressed_path
                    log.info("Compressão concluída.")
                else:
                    log.error(f"Falha na etapa de compressão: {comp_msg}")
                    # Opcional: retornar o estado anterior se a compressão falhar? Ou falhar tudo?
                    # Vamos falhar tudo se a compressão era obrigatória.
                    # Se OCR ainda for aplicado, podemos continuar do input original? Decidimos falhar.
                    if os.path.exists(str(compressed_path)): 
                        try:
                            compressed_path.unlink()
                        except:
                            pass
                    return False, None, f"Falha na compressão: {comp_msg}"

            # Etapa 2: OCR (se solicitado)
            if apply_ocr:
                ocr_output_path = Path(temp_dir) / "ocr_output.pdf"
                log.info("Aplicando OCR...")
                ocr_success, ocr_msg = self._apply_ocrmypdf(str(current_step_output), str(ocr_output_path), ocr_language)
                step_message = ocr_msg # Atualiza a mensagem da última etapa
                if ocr_success:
                    current_step_output = ocr_output_path
                    last_successful_output = ocr_output_path
                    log.info("OCR concluído.")
                else:
                    log.warning(f"Falha na etapa de OCR: {ocr_msg}. Continuando com o resultado anterior (se houver).")
                    # Se o OCR falhar, continuamos com o arquivo que entrou nesta etapa (last_successful_output)
                    current_step_output = last_successful_output # Reverte para o último sucesso
                    # Não retorna erro aqui, permite que o usuário receba pelo menos o resultado da compressão (se houve)
                    step_message = f"Compressão concluída, mas OCR falhou: {ocr_msg}"

            # Ler o resultado final
            if last_successful_output.exists() and last_successful_output.stat().st_size > 0:
                try:
                    with open(last_successful_output, 'rb') as f:
                        processed_bytes = f.read()
                    final_size_mb = len(processed_bytes) / 1024 / 1024
                    log.info(f"Processamento finalizado. Tamanho Final: {final_size_mb:.2f} MB. Mensagem: {step_message}")
                    return True, processed_bytes, step_message # Retorna a mensagem da última etapa principal
                except Exception as read_err:
                    log.error(f"Erro ao ler o arquivo processado final {last_successful_output}: {read_err}")
                    return False, None, "Erro ao ler o resultado final do processamento."
            else:
                log.error(f"Arquivo de resultado final ({last_successful_output}) não encontrado ou vazio após processamento.")
                return False, None, "Falha geral no processamento ou resultado final inválido."

    def image_to_pdf(self, image_files_bytes, output_pdf_path):
        """Converte uma lista de bytes de imagem em um único PDF."""
        if not image_files_bytes:
            return False, "Nenhuma imagem fornecida."

        valid_image_bytes_for_pdf = []
        processed_count = 0
        skipped_count = 0

        for idx, img_bytes in enumerate(image_files_bytes):
            try:
                with Image.open(io.BytesIO(img_bytes)) as img:
                    # Tenta converter para RGB (img2pdf lida melhor, evita problemas com paletas, RGBA, LA)
                    # Mantém PNG se for PNG, senão converte para JPEG para eficiência
                    output_format = "PNG" if img.format == 'PNG' else "JPEG"
                    # Se não for RGB/L (grayscale), converte
                    if img.mode not in ('RGB', 'L'):
                        log.debug(f"Convertendo imagem {idx+1} de modo {img.mode} para RGB.")
                        # Usar um fundo branco para transparência se converter de RGBA/LA/P
                        if img.mode in ('RGBA', 'LA', 'P'):
                            img = img.convert('RGB', palette=Image.ADAPTIVE, colors=256) # Tentativa com paleta
                            # Alternativa mais segura: criar fundo branco e colar
                            # background = Image.new("RGB", img.size, (255, 255, 255))
                            # background.paste(img, mask=img.split()[-1]) # Usa alpha como máscara
                            # img = background
                        else:
                            img = img.convert('RGB')

                    # Salva a imagem processada (possivelmente convertida) em bytes
                    byte_io = io.BytesIO()
                    img.save(byte_io, format=output_format, quality=85, optimize=True) # Adiciona quality/optimize
                    valid_image_bytes_for_pdf.append(byte_io.getvalue())
                    processed_count += 1

            except Exception as img_err:
                log.warning(f"Erro ao processar/converter imagem {idx+1}: {img_err}. Imagem pulada.")
                skipped_count += 1
                continue # Pula para a próxima imagem

        if not valid_image_bytes_for_pdf:
            log.error("Nenhuma imagem válida encontrada para converter para PDF.")
            return False, "Nenhuma imagem válida fornecida ou processada."

        try:
            # Converte os bytes das imagens válidas para PDF
            pdf_bytes = img2pdf.convert(valid_image_bytes_for_pdf)
            with open(output_pdf_path, "wb") as f:
                f.write(pdf_bytes)
            log.info(f"{processed_count} imagem(ns) convertida(s) para PDF: {output_pdf_path}. {skipped_count} imagem(ns) pulada(s).")
            return True, f"{processed_count} imagem(ns) convertida(s) para PDF. {skipped_count} pulada(s)."
        except Exception as e:
            log.exception(f"Erro ao criar PDF a partir das imagens: {str(e)}")
            return False, f"Erro ao gerar o arquivo PDF final: {e}"

    def pdf_to_docx(self, input_pdf_path, output_docx_path, apply_ocr=False, ocr_language='por'):
        """Converte PDF para DOCX, opcionalmente aplicando OCR antes."""
        if not fitz or not Document:
            log.error("PyMuPDF ou python-docx não disponíveis. Conversão PDF para DOCX impossível.")
            return False, "Bibliotecas necessárias para conversão não disponíveis."
            
        pdf_to_process = Path(input_pdf_path)
        temp_ocr_pdf_path = None
        ocr_applied_msg = ""

        if not pdf_to_process.exists():
            return False, "Arquivo PDF de entrada não encontrado."

        try:
            with tempfile.TemporaryDirectory() as temp_docx_dir_ctx:
                temp_docx_dir = Path(temp_docx_dir_ctx)

                # Aplicar OCR se solicitado e possível
                if apply_ocr:
                    if not self.ocrmypdf_installed:
                        log.warning("OCR solicitado para DOCX, mas OCRmyPDF não está disponível.")
                        ocr_applied_msg = " (OCR não aplicado: ferramenta indisponível)"
                    else:
                        log.info("Aplicando OCR antes da conversão para DOCX...")
                        temp_ocr_pdf_path = temp_docx_dir / "ocr_temp.pdf"
                        ocr_success, ocr_msg = self._apply_ocrmypdf(str(pdf_to_process), str(temp_ocr_pdf_path), ocr_language)
                        if ocr_success:
                            pdf_to_process = temp_ocr_pdf_path # Usa o PDF OCRizado
                            ocr_applied_msg = " (com OCR pré-aplicado)"
                            log.info("OCR aplicado com sucesso antes da conversão.")
                        else:
                            log.warning(f"Falha ao aplicar OCR antes da conversão para DOCX: {ocr_msg}. Procedendo com o PDF original.")
                            ocr_applied_msg = f" (tentativa de OCR falhou: {ocr_msg})"

                log.info(f"Iniciando conversão de {pdf_to_process.name} para DOCX{ocr_applied_msg}.")
                doc = fitz.open(str(pdf_to_process)) # PyMuPDF
                document = Document() # python-docx
                has_content = False

                # Processar cada página
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    # Extrair texto
                    text = page.get_text("text")
                    if text and text.strip():
                        document.add_paragraph(text)
                        has_content = True

                    # Extrair imagens
                    image_list = page.get_images(full=True)
                    for img_index, img_info in enumerate(image_list):
                        xref = img_info[0]
                        try:
                            base_image = doc.extract_image(xref)
                            image_bytes = base_image["image"]
                            image_ext = base_image["ext"]
                            # Salvar imagem temporariamente
                            temp_img_path = temp_docx_dir / f"p{page_num}_i{img_index}.{image_ext}"
                            with open(temp_img_path, "wb") as temp_img: temp_img.write(image_bytes)
                            # Adicionar imagem ao DOCX
                            try:
                                document.add_picture(str(temp_img_path), width=Inches(6.0)) # Largura padrão A4
                                has_content = True
                            except Exception as pic_err:
                                log.warning(f"Não foi possível adicionar imagem da pág {page_num+1} (idx {img_index}) ao DOCX: {pic_err}")
                        except Exception as img_extract_err:
                            log.warning(f"Não foi possível extrair imagem da pág {page_num+1} (idx {img_index}): {img_extract_err}")

                    # Adicionar quebra de página (exceto após a última)
                    if page_num < len(doc) - 1:
                        document.add_page_break()

                doc.close() # Fecha o documento PDF

                if not has_content:
                    log.warning(f"Conversão PDF->DOCX concluída, mas nenhum conteúdo (texto ou imagem) foi extraído para {output_docx_path}.")
                    # Não necessariamente um erro, pode ser um PDF em branco
                    message = "Conversão concluída, mas o DOCX pode estar vazio (nenhum conteúdo extraído)."
                else:
                    message = f"PDF convertido para DOCX com sucesso{ocr_applied_msg}."

                document.save(output_docx_path)
                log.info(f"Arquivo DOCX salvo em: {output_docx_path}")
                return True, message

        except Exception as e:
            log.exception(f"Erro durante a conversão PDF para DOCX: {str(e)}")
            return False, f"Erro inesperado durante a conversão: {e}"
        finally:
            # Limpeza adicional se necessário (embora temp dirs cuidem disso)
            pass

    def pdf_to_image(self, input_pdf_path, output_folder, image_format='png', dpi=150):
        """Converte cada página de um PDF em arquivos de imagem."""
        if not fitz:
            log.error("PyMuPDF não disponível. Conversão PDF para imagem impossível.")
            return None, "Biblioteca PyMuPDF necessária não disponível."
            
        generated_images = []
        doc = None
        output_path = Path(output_folder)

        if not Path(input_pdf_path).exists():
            return None, "Arquivo PDF de entrada não encontrado."

        try:
            # Cria pasta de saída se não existir
            output_path.mkdir(parents=True, exist_ok=True)

            doc = fitz.open(input_pdf_path)
            if not doc.page_count:
                return [], "PDF não contém páginas."

            log.info(f"Convertendo {len(doc)} página(s) PDF para {image_format.upper()} (DPI: {dpi}) em: {output_folder}")

            # Zoom baseado no DPI (72 é base)
            zoom = dpi / 72.0
            mat = fitz.Matrix(zoom, zoom)

            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                # Renderiza a página como imagem
                # alpha=False para evitar canal alfa (pode não ser suportado por todos os visualizadores/formatos)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                output_image_path = output_path / f"pagina_{page_num + 1}.{image_format.lower()}"

                # Salva a imagem
                pix.save(str(output_image_path))
                generated_images.append(str(output_image_path))
                log.debug(f"Salva imagem: {output_image_path.name}")

            log.info(f"{len(generated_images)} imagem(ns) gerada(s) com sucesso.")
            return generated_images, f"{len(generated_images)} imagem(ns) gerada(s)."

        except Exception as e:
            log.exception(f"Erro durante a conversão PDF para imagens: {str(e)}")
            return None, f"Erro inesperado na conversão: {e}"
        finally:
            if doc:
                try: doc.close()
                except: pass # Tenta fechar mesmo se houve erro antes

    def create_zip_from_files(self, file_paths, output_zip_path):
        """Cria um arquivo ZIP a partir de uma lista de caminhos de arquivo."""
        if not file_paths:
            return False, "Nenhum arquivo fornecido para compactar."

        try:
            log.info(f"Criando arquivo ZIP '{os.path.basename(output_zip_path)}' com {len(file_paths)} arquivo(s).")
            with zipfile.ZipFile(output_zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for file_path in file_paths:
                    if os.path.exists(file_path):
                        # Adiciona o arquivo ao ZIP usando apenas o nome base do arquivo
                        zipf.write(file_path, os.path.basename(file_path))
                    else:
                        log.warning(f"Arquivo não encontrado, pulando adição ao ZIP: {file_path}")

            if os.path.exists(output_zip_path) and os.path.getsize(output_zip_path) > 0:
                log.info(f"Arquivo ZIP criado com sucesso: {output_zip_path}")
                return True, "Arquivo ZIP criado com sucesso."
            elif not os.path.exists(output_zip_path):
                # Isso não deveria acontecer se o zipfile não deu erro, mas checar nunca é demais
                log.error("Arquivo ZIP não foi criado após a operação.")
                return False, "Falha ao criar o arquivo ZIP (não encontrado)."
            else:
                # ZIP foi criado mas está vazio (talvez nenhum arquivo de entrada válido?)
                log.warning(f"Arquivo ZIP criado ({output_zip_path}), mas está vazio.")
                return True, "Arquivo ZIP criado, mas pode estar vazio (verifique arquivos de entrada)."


        except Exception as e:
            log.exception(f"Erro ao criar arquivo ZIP: {str(e)}")
            # Tentar remover zip parcial
            if os.path.exists(output_zip_path):
                try: os.unlink(output_zip_path)
                except OSError: pass
            return False, f"Erro ao criar o arquivo ZIP: {e}"

    def document_to_pdf(self, input_doc_path, output_pdf_path):
        """Converte DOC, DOCX, ODT, TXT, etc. para PDF usando LibreOffice."""
        if not self.libreoffice_path:
            log.error("LibreOffice não encontrado. Conversão de documento indisponível.")
            return False, "LibreOffice não está disponível no servidor."

        if not Path(input_doc_path).exists():
            return False, "Arquivo de documento de entrada não encontrado."

        output_dir = Path(output_pdf_path).parent
        # Garante que o diretório de saída exista
        output_dir.mkdir(parents=True, exist_ok=True)

        # O LibreOffice pode gerar um nome diferente, então precisamos encontrar/renomear
        input_basename_no_ext = Path(input_doc_path).stem
        # Caminho esperado se LO apenas mudar a extensão
        expected_lo_output_simple = output_dir / f"{input_basename_no_ext}.pdf"

        # Limpar saídas anteriores potenciais para evitar confusão
        if Path(output_pdf_path).exists(): Path(output_pdf_path).unlink()
        if expected_lo_output_simple.exists() and expected_lo_output_simple != Path(output_pdf_path):
            expected_lo_output_simple.unlink()

        # Comando para conversão
        command = [
            self.libreoffice_path, '--headless', '--invisible', # invisible pode ajudar em alguns casos
            '--convert-to', 'pdf', '--outdir', str(output_dir),
            input_doc_path
        ]

        # Usar helper para rodar LO
        success, result_msg_or_proc = self._run_subprocess(command, "Conversão LibreOffice", timeout=300) # Timeout LO

        if not success:
            return False, result_msg_or_proc # Mensagem de erro do helper

        # Verificar se a saída esperada (ou a final) foi criada
        final_pdf_found = None
        if Path(output_pdf_path).exists() and Path(output_pdf_path).stat().st_size > 0:
            final_pdf_found = Path(output_pdf_path)
        elif expected_lo_output_simple.exists() and expected_lo_output_simple.stat().st_size > 0:
            # Renomear para o nome desejado
            try:
                shutil.move(str(expected_lo_output_simple), output_pdf_path)
                log.info(f"Arquivo LibreOffice renomeado para: {os.path.basename(output_pdf_path)}")
                final_pdf_found = Path(output_pdf_path)
            except Exception as move_err:
                log.error(f"LibreOffice gerou {expected_lo_output_simple.name}, mas falhou ao renomear para {os.path.basename(output_pdf_path)}: {move_err}")
                return False, "Erro ao renomear o PDF gerado pelo LibreOffice."
        # Adicionar verificação para outros nomes possíveis se LO for imprevisível?

        if final_pdf_found:
            log.info(f"Documento convertido para PDF com sucesso: {output_pdf_path}")
            return True, "Documento convertido para PDF com sucesso."
        else:
            log.error(f"LibreOffice executado, mas o arquivo PDF final não foi encontrado ou está vazio em {output_dir}.")
            # Logar stdout/stderr do processo LO para debug
            if isinstance(result_msg_or_proc, subprocess.CompletedProcess):
                log.error(f"Stderr LO: {result_msg_or_proc.stderr.strip()}\nStdout LO: {result_msg_or_proc.stdout.strip()}")
            return False, "Falha na criação do arquivo PDF pelo LibreOffice."


    def merge_pdfs(self, pdf_byte_streams):
        """Junta múltiplos streams de bytes de PDF em um único stream de bytes."""
        if not pdf_byte_streams or len(pdf_byte_streams) < 2:
            return False, None, "Pelo menos dois PDFs são necessários para a junção."

        merged_writer = PdfWriter()
        log.info(f"Iniciando junção de {len(pdf_byte_streams)} PDF(s).")
        valid_pdfs_merged = 0

        try:
            for idx, pdf_bytes in enumerate(pdf_byte_streams):
                try:
                    # Usar BytesIO para ler os bytes como um arquivo
                    pdf_stream = io.BytesIO(pdf_bytes)
                    reader = PdfReader(pdf_stream)

                    if not reader.pages:
                        log.warning(f"PDF {idx+1} está vazio ou corrompido, pulando.")
                        continue

                    # Adicionar todas as páginas do PDF atual
                    for page in reader.pages:
                        merged_writer.add_page(page)

                    valid_pdfs_merged += 1
                    log.debug(f"Adicionado PDF {idx+1} ({len(reader.pages)} páginas).")

                except Exception as read_err:
                    # Logar erro específico na leitura/processamento de um PDF, mas continuar
                    log.error(f"Erro ao processar PDF {idx+1}: {read_err}. Pulando este PDF.")
                    # Considerar se deve falhar tudo ou apenas pular o PDF problemático
                    # Aqui estamos pulando.

            # Verificar se algum PDF válido foi adicionado
            if valid_pdfs_merged == 0:
                return False, None, "Nenhum PDF válido encontrado para juntar."
            if valid_pdfs_merged < len(pdf_byte_streams):
                merge_message_suffix = f" ({valid_pdfs_merged} de {len(pdf_byte_streams)} PDFs juntados com sucesso)."
            else:
                merge_message_suffix = "."

            # Escrever o resultado em um stream de bytes na memória
            output_stream = io.BytesIO()
            merged_writer.write(output_stream)
            # merged_writer.close() # PdfWriter não tem close() explícito necessário como arquivos

            # Obter os bytes resultantes
            output_stream.seek(0)
            result_bytes = output_stream.getvalue()

            if not result_bytes:
                # Isso seria estranho se valid_pdfs_merged > 0
                log.error("Junção de PDFs resultou em um arquivo vazio.")
                return False, None, "Erro inesperado: Junção resultou em arquivo vazio."

            log.info(f"Junção de PDFs concluída com sucesso{merge_message_suffix}")
            return True, result_bytes, f"PDFs juntados com sucesso{merge_message_suffix}"

        except Exception as e:
            error_msg = f"Erro inesperado durante a junção de PDFs: {e}"
            log.exception(error_msg)
            return False, None, error_msg
        finally:
            # Garantir que o writer seja fechado (embora não seja estritamente necessário para BytesIO?)
            pass # PdfWriter não tem close()